import streamlit as st, pydicom, io, os, pytesseract
from PIL import Image
from docx import Document
from openai import OpenAI

# Azmed etiket → Türkçe karşılık
PHRASES = {
    "PLR.EFF": "Plevral efüzyon",
    "FRAC.":   "Kırık",
    "DISLOC.": "Dislokasyon",
    "NODULE":  "Nodül",
    "DOUBT":   "Şüpheli bulgu"
}

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def dcm2png(raw):
    ds = pydicom.dcmread(io.BytesIO(raw))
    return Image.fromarray(ds.pixel_array).convert("RGB")

def find_tags(img):
    txt = pytesseract.image_to_string(img)
    return [k for k in PHRASES if k in txt]

def make_draft(tags):
    prompt = (f"Etiketler: {', '.join(tags) or 'yok'}. "
              "Sağlık Bakanlığı Teleradyoloji formatında "
              "(TEKNİK / KARŞILAŞTIRMA / BULGULAR / SONUÇ) "
              "Türkçe bir radyoloji raporu taslağı yaz. "
              "Marka veya yazılım adı geçmesin.")
    r = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role":"user","content":prompt}],
        temperature=0)
    return r.choices[0].message.content

def to_docx(text):
    doc = Document("sb_template.docx")
    heads = {h: "" for h in ["TEKNİK","KARŞILAŞTIRMA","BULGULAR","SONUÇ"]}
    for block in text.split("\n\n"):
        for h in heads:
            if block.upper().startswith(h):
                heads[h] = block.split("\n",1)[1] if "\n" in block else ""
    for p in doc.paragraphs:
        if p.text in heads:
            p.clear(); p.add_run(heads[p.text])
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ---------- UI ----------
st.set_page_config("TemplAID", layout="wide")
st.title("📑 TemplAID — Taslak Radyoloji Raporu")

file = st.file_uploader("Azmed DICOM dosyasını yükleyin", type=["dcm"])
if file:
    img = dcm2png(file.read())
    st.image(img, width=350)
    tags = find_tags(img)
    st.success(f"Etiketler: {', '.join(tags) or 'yok'}")
    
    if st.button("Taslak Rapor Oluştur"):
        draft = make_draft(tags)
        edited = st.text_area("📝 Düzenle (gerekirse değiştirin)", draft, height=400)
        st.download_button(
            "DOCX indir", to_docx(edited), "taslak_rapor.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
